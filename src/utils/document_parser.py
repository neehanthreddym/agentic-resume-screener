"""
Document parsing utilities for the AI Resume Screener.

Proven implementations from notebooks/01_document_parsing.ipynb.

Key decisions:
- PDF: PyMuPDF (fitz) block-level extraction — handles multi-column layouts correctly.
  pdfplumber's default extract_text() reads by Y-coordinate and scrambles two-column resumes.
- DOCX: XML traversal (iter_block_items) — walks doc.element.body children in order,
  interleaving paragraphs and tables exactly as they appear visually. The naive two-pass
  approach (doc.paragraphs then doc.tables) destroys reading order and silently skips
  table content in resumes like Olivia Harper's which use invisible tables for layout.
- TXT: Standard file read — no special handling needed.
"""
from pathlib import Path
from typing import Union

import docx
import fitz  # PyMuPDF
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def _iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.

    This solves the python-docx hierarchy issue by reading the raw XML sequence
    instead of iterating doc.paragraphs and doc.tables as two separate lists.
    Critical for resumes like Maya Srinivasan's where the CORE SKILLS table sits
    between PROFESSIONAL SUMMARY and EDUCATION — not at the end of the document.
    """
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"Unsupported parent type: {type(parent)}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parse_pdf(file_path: Union[str, Path]) -> str:
    """
    Extract text from a PDF using PyMuPDF block-level extraction.
    Preserves column reading order for multi-column layouts.

    Args:
        file_path: Path to the PDF file.

    Returns:
        Extracted text as a string. Returns empty string on failure.
    """
    text = ""
    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            # get_text("blocks") returns a list of (x0, y0, x1, y1, "text", ...) tuples.
            # PyMuPDF returns blocks in logical reading order, not raw Y-stripe order.
            blocks = page.get_text("blocks")
            for block in blocks:
                text += block[4]  # index 4 is the text string
        return text.strip()
    except Exception as e:
        print(f"[Parser] Error reading PDF {file_path}: {e}")
        return ""


def parse_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX file using XML traversal to preserve document hierarchy.

    Walks doc.element.body children in order, interleaving paragraphs and tables
    exactly as they appear in the document — NOT paragraphs-first, tables-second.
    Handles merged cells by deduplicating cell text within a row.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a string. Returns empty string on failure.
    """
    try:
        doc = docx.Document(str(file_path))
        full_text = []

        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                if block.text.strip():
                    full_text.append(block.text.strip())

            elif isinstance(block, Table):
                for row in block.rows:
                    row_data = []
                    for cell in row.cells:
                        clean_cell = cell.text.strip()
                        # Skip empty cells and deduplicate merged cells
                        if clean_cell and clean_cell not in row_data:
                            row_data.append(clean_cell)
                    if row_data:
                        full_text.append("\n".join(row_data))

        return "\n".join(full_text)

    except Exception as e:
        print(f"[Parser] Error reading DOCX {file_path}: {e}")
        return ""


def parse_txt(file_path: Union[str, Path]) -> str:
    """
    Read a plain text file.

    Args:
        file_path: Path to the TXT file.

    Returns:
        File contents as a string. Returns empty string on failure.
    """
    try:
        with open(str(file_path), "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception as e:
        print(f"[Parser] Error reading TXT {file_path}: {e}")
        return ""


def parse_document(file_path: Union[str, Path]) -> str:
    """
    Route a document to the correct parser based on file extension.
    Supports PDF, DOCX, and TXT.

    Args:
        file_path: Path to the document.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file extension is not supported.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(path)
    elif suffix == ".docx":
        return parse_docx(path)
    elif suffix == ".txt":
        return parse_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type: '{suffix}'. Supported types: .pdf, .docx, .txt"
        )
